"""
Case generation: extract page context → CaseBuilder LLM.

Script testcase LangGraph (planner/execute/evaluate) lives in `services.script.graph`.
"""

import base64
import io
import json
import re
from typing import Any
from langchain.messages import HumanMessage, SystemMessage

from ..llm import llm_vision, llm_extraction
from ..extraction import SourceExtraction
from .prompt import CaseBuilder_SystemPrompt, TestcaseDraft_SystemPrompt
from .state import debug_case_node_enter


def _guess_mime_from_filename(name: str) -> str:
    n = (name or "").lower()
    if n.endswith(".png"):
        return "image/png"
    if n.endswith((".jpg", ".jpeg")):
        return "image/jpeg"
    if n.endswith(".webp"):
        return "image/webp"
    if n.endswith(".gif"):
        return "image/gif"
    if n.endswith(".pdf"):
        return "application/pdf"
    if n.endswith(".txt"):
        return "text/plain"
    if n.endswith(".csv"):
        return "text/csv"
    if n.endswith(".docx"):
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if n.endswith(".xlsx"):
        return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    if n.endswith(".xls"):
        return "application/vnd.ms-excel"
    return "application/octet-stream"


def _extract_pdf_text(data: bytes) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        parts: list[str] = []
        for page in reader.pages[:30]:
            t = page.extract_text()
            if t:
                parts.append(t)
        return "\n\n".join(parts)[:14000]
    except Exception:
        return ""


_DOC_TABULAR_MAX_CHARS = 16000


def _extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document

        doc = Document(io.BytesIO(data))
        parts: list[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells = [(c.text or "").strip() for c in row.cells]
                if any(cells):
                    parts.append("\t".join(cells))
        return "\n\n".join(parts)[:_DOC_TABULAR_MAX_CHARS]
    except Exception:
        return ""


def _extract_xlsx_text(data: bytes) -> str:
    wb = None
    try:
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts: list[str] = []
        char_budget = _DOC_TABULAR_MAX_CHARS
        for sheet in wb.worksheets:
            parts.append(f"## Sheet: {sheet.title}")
            for ri, row in enumerate(sheet.iter_rows(values_only=True)):
                if ri > 500:
                    break
                line = "\t".join(
                    "" if v is None else str(v).strip() for v in row
                ).strip()
                if line:
                    parts.append(line)
                chunk = "\n".join(parts)
                if len(chunk) >= char_budget:
                    return chunk[:char_budget]
        return "\n".join(parts)[:_DOC_TABULAR_MAX_CHARS]
    except Exception:
        return ""
    finally:
        if wb is not None:
            try:
                wb.close()
            except Exception:
                pass


def _extract_xls_text(data: bytes) -> str:
    try:
        import xlrd

        book = xlrd.open_workbook(file_contents=data)
        parts: list[str] = []
        char_budget = _DOC_TABULAR_MAX_CHARS
        for si in range(book.nsheets):
            sh = book.sheet_by_index(si)
            parts.append(f"## Sheet: {sh.name}")
            for r in range(min(sh.nrows, 500)):
                vals: list[str] = []
                for c in range(sh.ncols):
                    try:
                        vals.append(str(sh.cell_value(r, c)).strip())
                    except Exception:
                        vals.append("")
                line = "\t".join(vals).strip()
                if line:
                    parts.append(line)
                if len("\n".join(parts)) >= char_budget:
                    return "\n".join(parts)[:char_budget]
        return "\n".join(parts)[:_DOC_TABULAR_MAX_CHARS]
    except Exception:
        return ""


def _normalize_attachment_entries(raw: list) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    for i, item in enumerate(raw or []):
        if not isinstance(item, dict):
            continue
        name = str(item.get("file_name") or f"file_{i + 1}").strip() or f"file_{i + 1}"
        b64 = item.get("file_content_base64") or ""
        detail = str(item.get("file_detail") or "").strip()
        mime = str(item.get("mime_type") or "").strip() or _guess_mime_from_filename(name)
        out.append(
            {
                "file_name": name,
                "file_detail": detail,
                "mime_type": mime,
                "file_content_base64": b64,
            }
        )
    return out


def _attachment_meta_only(entries: list[dict[str, Any]]) -> list[dict[str, str]]:
    return [
        {
            "file_name": str(e.get("file_name", "")),
            "file_detail": str(e.get("file_detail", "")),
            "mime_type": str(e.get("mime_type", "")),
        }
        for e in entries
    ]


def debug_case_log_attachments_summary(entries: list[dict[str, Any]]) -> None:
    """Stdout: per-file payload size and text extraction stats (DOCX/PDF/Excel/txt). Missing base64 is explicit."""
    if not entries:
        print("[case-graph] test_attachments: (none)", flush=True)
        return
    print(f"[case-graph] test_attachments: {len(entries)} file(s)", flush=True)
    for e in entries:
        name = str(e.get("file_name") or "(unnamed)")
        mime = str(e.get("mime_type") or "").strip() or _guess_mime_from_filename(name)
        b64 = str(e.get("file_content_base64") or "").strip()
        if not b64:
            print(
                f"  · {name} | {mime}\n"
                f"    ⚠ file_content_base64 empty — ไฟล์นี้ถูกข้าม ไม่มีการอ่าน DOCX/PDF/รูป (ฝั่ง client ต้องส่ง base64 จริง)",
                flush=True,
            )
            continue
        try:
            raw = base64.b64decode(b64, validate=False)
        except Exception as exc:
            print(
                f"  · {name} | {mime}\n    ⚠ base64 decode failed: {exc}",
                flush=True,
            )
            continue
        n = len(raw)
        nlower = name.lower()
        extra = ""
        if mime.startswith("image/") or nlower.endswith((".png", ".jpg", ".jpeg", ".webp", ".gif")):
            kind = "image"
            extra = " → sent as image_url to multimodal model"
        elif mime == "text/plain" or nlower.endswith(".txt"):
            kind = "txt"
            t = raw.decode("utf-8", errors="replace")
            extra = f" | text_chars={len(t.strip())}"
        elif mime in ("text/csv", "application/csv") or nlower.endswith(".csv"):
            kind = "csv"
            t = raw.decode("utf-8", errors="replace")
            extra = f" | text_chars={len(t.strip())}"
        elif "wordprocessingml" in mime.lower() or nlower.endswith(".docx"):
            kind = "docx"
            t = _extract_docx_text(raw)
            nc = len(t.strip())
            extra = f" | extracted_text_chars={nc}" + (" ✓" if nc else " ⚠ (no text extracted)")
        elif "pdf" in mime.lower() or nlower.endswith(".pdf"):
            kind = "pdf"
            t = _extract_pdf_text(raw)
            nc = len(t.strip())
            extra = f" | extracted_text_chars={nc}" + (" ✓" if nc else " ⚠ (no text extracted)")
        elif (
            "spreadsheetml" in mime.lower()
            or mime == "application/vnd.ms-excel"
            or nlower.endswith((".xlsx", ".xls"))
        ):
            kind = "excel"
            t = _extract_xls_text(raw) if nlower.endswith(".xls") else _extract_xlsx_text(raw)
            nc = len(t.strip())
            extra = f" | extracted_text_chars={nc}" + (" ✓" if nc else " ⚠ (no text extracted)")
        else:
            kind = "other"
            extra = " → metadata only in prompt (unsupported type for text extract)"
        print(f"  · {name} | {kind} | payload_bytes={n}{extra}", flush=True)


def _append_attachment_blocks_to_content(
    content_blocks: list[dict[str, Any]],
    entries: list[dict[str, Any]],
) -> None:
    for e in entries:
        name = str(e.get("file_name", "file"))
        detail = str(e.get("file_detail", "")).strip()
        mime = str(e.get("mime_type", "")).strip() or _guess_mime_from_filename(name)
        b64 = str(e.get("file_content_base64") or "").strip()
        if not b64:
            continue
        try:
            raw = base64.b64decode(b64, validate=False)
        except Exception:
            continue

        caption = (
            f"### User-attached file: {name}\n"
            f"**User description (what this file is for):** {detail or '(none)'}\n\n"
        )

        is_image = mime.startswith("image/") or name.lower().endswith(
            (".png", ".jpg", ".jpeg", ".webp", ".gif")
        )
        if is_image:
            img_mime = mime if mime.startswith("image/") else "image/png"
            content_blocks.append({"type": "text", "text": caption})
            content_blocks.append(
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:{img_mime};base64,{b64}"},
                }
            )
            continue

        nlower = name.lower()

        if mime == "text/plain" or nlower.endswith(".txt"):
            try:
                txt = raw.decode("utf-8", errors="replace")[:16000]
            except Exception:
                txt = ""
            content_blocks.append(
                {"type": "text", "text": caption + f"```text\n{txt}\n```"}
            )
            continue

        if mime in ("text/csv", "application/csv") or nlower.endswith(".csv"):
            try:
                txt = raw.decode("utf-8", errors="replace")[:16000]
            except Exception:
                txt = ""
            content_blocks.append(
                {
                    "type": "text",
                    "text": caption + f"```text\n{txt}\n```",
                }
            )
            continue

        if "wordprocessingml" in mime.lower() or nlower.endswith(".docx"):
            docx_text = _extract_docx_text(raw)
            if docx_text.strip():
                content_blocks.append(
                    {"type": "text", "text": caption + f"```text\n{docx_text}\n```"}
                )
            else:
                content_blocks.append(
                    {
                        "type": "text",
                        "text": caption
                        + "(DOCX text could not be extracted; use the user description and filename.)\n",
                    }
                )
            continue

        if (
            "spreadsheetml" in mime.lower()
            or mime == "application/vnd.ms-excel"
            or nlower.endswith((".xlsx", ".xls"))
        ):
            if nlower.endswith(".xls"):
                sheet_text = _extract_xls_text(raw)
            else:
                sheet_text = _extract_xlsx_text(raw)
            if sheet_text.strip():
                content_blocks.append(
                    {"type": "text", "text": caption + f"```text\n{sheet_text}\n```"}
                )
            else:
                content_blocks.append(
                    {
                        "type": "text",
                        "text": caption
                        + "(Spreadsheet text could not be extracted; use the user description and filename.)\n",
                    }
                )
            continue

        if "pdf" in mime.lower() or name.lower().endswith(".pdf"):
            pdf_text = _extract_pdf_text(raw)
            if pdf_text.strip():
                content_blocks.append(
                    {"type": "text", "text": caption + f"```\n{pdf_text}\n```"}
                )
            else:
                content_blocks.append(
                    {
                        "type": "text",
                        "text": caption
                        + "(PDF text could not be extracted; use the user description and filename.)\n",
                    }
                )
            continue

        content_blocks.append(
            {
                "type": "text",
                "text": caption + f"(Unsupported MIME `{mime}`; metadata only.)\n",
            }
        )


def _normalize_case_builder_display_name(raw_name: str, description: str) -> str:
    base = (raw_name or "").strip()
    base = re.sub(r"(?i)^TC\d+[_\-\s]*", "", base).strip()
    base = re.sub(r"(?i)^TC[_\s]+", "", base).strip()
    if not base:
        base = (description or "").strip()
    base = re.sub(r"\s+", " ", base)
    if len(base) > 120:
        base = base[:117].rstrip() + "..."
    return base or "Test case"


def _normalize_single_testcase_enhance_output(name: str, description: str) -> tuple[str, str]:
    n = _normalize_case_builder_display_name(name or "", "")
    d = (description or "").strip()
    d = re.sub(r"\s+", " ", d)
    if len(d) > 8000:
        d = d[:7997] + "..."
    return n, d


class CaseAgentEntryPoint:
    """Entry points for AI testcase generation (page extract + CaseBuilder LLM)."""

    @staticmethod
    def _parse_case_generate_request(data: dict) -> tuple[str, str, str, list]:
        """test_name, test_url, test_spec, raw_attachment_list."""
        base = data if isinstance(data, dict) else {}
        td = base.get("test_detail")
        td = td if isinstance(td, dict) else {}
        test_name = td.get("test_name") or base.get("test_name") or "Untitled Test"
        test_url = td.get("test_url") or base.get("test_url") or ""
        test_spec = td.get("test_spec") or base.get("test_spec") or ""
        raw = td.get("test_attachments") or base.get("test_attachments") or []
        return str(test_name), str(test_url), str(test_spec), raw if isinstance(raw, list) else []

    @staticmethod
    def extract_attachments(raw_list: list) -> tuple[list[dict[str, Any]], list[dict[str, str]]]:
        """Normalize uploads and build metadata-only rows for the text prompt."""
        entries = _normalize_attachment_entries(raw_list)
        return entries, _attachment_meta_only(entries)

    @staticmethod
    async def extract_url(test_url: str, test_name: str, test_spec: str) -> dict:
        """Run `SourceExtraction` for the target URL (HTML, DOM summary, screenshot, etc.)."""
        result = await SourceExtraction.fetch(
            test_url,
            test_name=test_name,
            test_spec=test_spec,
        )
        return result if isinstance(result, dict) else {}

    @staticmethod
    def merge_case_builder_input(
        test_name: str,
        test_url: str,
        test_spec: str,
        attachment_meta: list[dict[str, str]],
        attachment_entries: list[dict[str, Any]],
        extraction: dict,
    ) -> tuple[str, list[dict[str, Any]]]:
        """One string prefix plus multimodal `content_blocks` (page screenshot + decoded files)."""
        prompt_context = extraction.get("prompt_context", "")
        cleaned_html = extraction.get("cleaned_html", "")
        main_elements = extraction.get("main_elements", {})
        grouped_elements = extraction.get("grouped_elements", {})
        screenshot_b64 = extraction.get("page_screenshot", "")

        att_block = ""
        if attachment_entries:
            att_block = (
                "## USER-ATTACHED FILES (mandatory)\n"
                "The user sent **one or more** reference files. Each row in USER_ATTACHMENT_METADATA includes "
                "**`file_detail`** — that text is the user’s intent for that file (requirement, screen area, flow to cover).\n"
                "- **Must** use attachments together with page extraction: add or adjust testcases when specs/mockups/PDFs "
                "imply coverage the live page HTML alone does not state.\n"
                "- **Must** respect **`file_detail`** per file; if two files differ, split or label scenarios clearly.\n"
                "- Images: treat as UI/spec evidence. PDF/TXT/CSV and **DOCX/Excel (converted to text server-side)**: extracted text "
                "appears in the message below; if extraction failed or **`file_content_base64` was empty**, rely on **`file_detail`** + filename.\n\n"
            )

        user_text = (
            f"test_name: {test_name}\n"
            f"test_url: {test_url}\n"
            f"test_spec: {test_spec}\n\n"
            f"{att_block}"
            "## USER_ATTACHMENT_METADATA (names + user descriptions; binary content is attached separately)\n"
            f"{json.dumps(attachment_meta, ensure_ascii=False, indent=2)}\n\n"
            "Use this extraction context to generate testcases.\n"
            "When USER_ATTACHMENT_METADATA is non-empty, incorporate those files and the user's per-file **`file_detail`** into relevant testcases.\n"
            f"{prompt_context}\n\n"
            f"main_elements_json:\n{json.dumps(main_elements, ensure_ascii=False, indent=2)}\n\n"
            f"grouped_elements_json:\n{json.dumps(grouped_elements, ensure_ascii=False, indent=2)[:16000]}\n\n"
            f"cleaned_html:\n{str(cleaned_html)[:6000]}"
        )

        content_blocks: list[dict[str, Any]] = [{"type": "text", "text": user_text}]
        if screenshot_b64:
            content_blocks.append(
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:image/png;base64,{screenshot_b64}"},
                }
            )
        _append_attachment_blocks_to_content(content_blocks, attachment_entries)
        return user_text, content_blocks

    @staticmethod
    async def call_case_builder_llm(content_blocks: list[dict[str, Any]]) -> str:
        """System prompt + multimodal human message → raw model string."""
        messages: list[Any] = [
            SystemMessage(content=CaseBuilder_SystemPrompt),
            HumanMessage(content=content_blocks),
        ]
        vision = any(b.get("type") == "image_url" for b in content_blocks)
        model = llm_vision if vision else llm_extraction
        response = await model.ainvoke(messages)
        if isinstance(response.content, str):
            return response.content.strip()
        return str(response.content).strip()

    @staticmethod
    def simplified_testcases_from_llm_json(parsed: dict) -> list[dict[str, Any]]:
        rows: list[dict[str, Any]] = []
        for tc in parsed.get("testcase", []) if isinstance(parsed, dict) else []:
            if not isinstance(tc, dict):
                continue
            desc = tc.get("testcase_description", "") or ""
            name = _normalize_case_builder_display_name(tc.get("testcase_name", "") or "", desc)
            cat_raw = tc.get("testcase_category", "") or ""
            cat = str(cat_raw).strip() or "functional"
            rows.append(
                {
                    "testcase_name": name,
                    "testcase_description": desc,
                    "testcase_category": cat,
                    "testcase_priority": tc.get("testcase_priority", "Medium"),
                }
            )
        return rows

    @staticmethod
    async def case_generate(data: dict) -> dict:
        test_name, test_url, test_spec, raw_attachments = CaseAgentEntryPoint._parse_case_generate_request(
            data if isinstance(data, dict) else {}
        )
        attachment_entries, attachment_meta = CaseAgentEntryPoint.extract_attachments(raw_attachments)
        extraction = await CaseAgentEntryPoint.extract_url(test_url, test_name, test_spec)
        _, content_blocks = CaseAgentEntryPoint.merge_case_builder_input(
            test_name,
            test_url,
            test_spec,
            attachment_meta,
            attachment_entries,
            extraction,
        )

        raw = await CaseAgentEntryPoint.call_case_builder_llm(content_blocks)
        simplified_cases = CaseAgentEntryPoint.simplified_testcases_from_llm_json(CaseAgentEntryPoint._parse_json(raw))

        grouped_elements = extraction.get("grouped_elements", {})
        page_extraction_str = json.dumps(grouped_elements, ensure_ascii=False, separators=(",", ":"))
        ex_errors = extraction.get("errors")
        if not isinstance(ex_errors, list):
            ex_errors = []

        return {
            "test_name": test_name,
            "test_url": test_url,
            "test_spec": test_spec,
            "test_attachments": attachment_meta,
            "page_extraction": {
                "ok": bool(extraction.get("ok", False)),
                "errors": ex_errors,
                "page_screenshot": extraction.get("page_screenshot", ""),
                "page_extraction": page_extraction_str,
            },
            "testcase": simplified_cases,
        }

    @staticmethod
    async def case_draft(data: dict) -> dict:
        if not isinstance(data, dict):
            data = {}
        current_name = str(data.get("current_testcase_name") or "").strip()
        current_desc = str(data.get("current_testcase_description") or "").strip()
        test_detail = data.get("test_detail") if isinstance(data.get("test_detail"), dict) else {}
        test_name = str(test_detail.get("test_name") or "").strip()
        test_url = str(test_detail.get("test_url") or "").strip()
        test_spec = str(test_detail.get("test_spec") or "").strip()
        raw_att = test_detail.get("test_attachments")
        att_list = raw_att if isinstance(raw_att, list) else []
        attachment_entries = _normalize_attachment_entries(att_list)
        attachment_meta = _attachment_meta_only(attachment_entries)
        debug_case_log_attachments_summary(attachment_entries)

        user_lines = [
            "## User draft (single testcase)",
            f"current_testcase_name: {current_name}",
            f"current_testcase_description: {current_desc}",
            "",
            "## Test context (may be empty)",
            f"test_name: {test_name}",
            f"test_url: {test_url}",
            f"test_spec: {test_spec}",
        ]
        user_text = "\n".join(user_lines)

        content_blocks: list[dict[str, Any]] = [{"type": "text", "text": user_text}]
        screenshot_b64 = ""

        if test_url:
            extraction = await SourceExtraction.fetch(
                test_url,
                test_name=test_name or "Test",
                test_spec=test_spec,
            )
            prompt_context = extraction.get("prompt_context", "")
            grouped_elements = extraction.get("grouped_elements", {})
            screenshot_b64 = extraction.get("page_screenshot", "") or ""

            extra = (
                "\n\n## Page extraction (for wording only)\n"
                f"{prompt_context}\n\n"
                f"grouped_elements_json (truncated):\n{json.dumps(grouped_elements, ensure_ascii=False, indent=2)[:12000]}\n"
            )
            content_blocks[0]["text"] = user_text + extra

        if attachment_meta:
            content_blocks[0]["text"] = content_blocks[0]["text"] + (
                "\n\n## USER_ATTACHMENT_METADATA\n"
                f"{json.dumps(attachment_meta, ensure_ascii=False, indent=2)}\n"
                "When polishing, apply each file’s **`file_detail`** together with any extracted text (PDF/TXT/CSV/DOCX/Excel) or image blocks below.\n"
            )

        _append_attachment_blocks_to_content(content_blocks, attachment_entries)

        messages: list[Any] = [SystemMessage(content=TestcaseDraft_SystemPrompt)]
        if screenshot_b64:
            content_blocks.append(
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:image/png;base64,{screenshot_b64}"},
                }
            )

        messages.append(HumanMessage(content=content_blocks))

        has_images = any(b.get("type") == "image_url" for b in content_blocks)
        model = llm_vision if has_images else llm_extraction
        response = await model.ainvoke(messages)
        raw = response.content.strip() if isinstance(response.content, str) else str(response.content)

        parsed = CaseAgentEntryPoint._parse_json(raw)
        if not isinstance(parsed, dict):
            parsed = {}
        out_name = parsed.get("testcase_name") or current_name
        out_desc = parsed.get("testcase_description") or current_desc
        out_name, out_desc = _normalize_single_testcase_enhance_output(str(out_name), str(out_desc))

        out: dict[str, Any] = {
            "testcase_name": out_name,
            "testcase_description": out_desc,
        }
        second = parsed.get("suggested_second_testcase")
        if isinstance(second, dict):
            s_name = (second.get("testcase_name") or "").strip()
            s_desc = (second.get("testcase_description") or "").strip()
            if s_name or s_desc:
                out["suggested_second_testcase"] = {
                    "testcase_name": s_name,
                    "testcase_description": s_desc,
                }

        return out

    @staticmethod
    def _parse_json(raw: str) -> dict:
        content = (raw or "").strip()
        if content.startswith("```"):
            content = content.split("\n", 1)[1] if "\n" in content else content[3:]
            if content.endswith("```"):
                content = content[:-3]
            content = content.strip()
        try:
            return json.loads(content)
        except Exception:
            try:
                start = content.find("{")
                end = content.rfind("}")
                if start != -1 and end != -1 and end > start:
                    return json.loads(content[start : end + 1])
            except Exception:
                pass
            return {"testcase": []}
